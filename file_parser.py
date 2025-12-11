# file_parser.py

import io
from fastapi import UploadFile
from typing import Optional
import docx
import fitz


def parse_docx(file_content: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = '\n'.join([p.text for p in doc.paragraphs])
        if not text.strip():
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text for cell in row.cells])
                    text += row_text + '\n'
        return text.strip()
    except Exception as e:
        raise Exception(f"解析 Word 文件失败: {e}")


def parse_pdf(file_content: bytes) -> str:
    try:
        document = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in document:
            text += page.get_text() + "\n"
        print(text)
        return text.strip()
    except Exception as e:
        raise Exception(f"解析 PDF 文件失败: {e}")


async def parse_uploaded_file(file: UploadFile) -> str:
    await file.seek(0)
    content = await file.read()

    mime_type = file.content_type

    if mime_type == "application/pdf":
        return parse_pdf(content)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx(content)
    elif mime_type.startswith('text/'):
        return content.decode('utf-8')
    else:
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            raise ValueError(f"不支持或无法识别的文件类型: {mime_type} ({file.filename})")